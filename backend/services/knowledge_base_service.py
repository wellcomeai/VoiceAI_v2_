import os
import uuid
from typing import List, Dict, Any, Optional
import json
import asyncio
from datetime import datetime
import tiktoken
import io

from sqlalchemy.orm import Session
from fastapi import UploadFile, HTTPException, status

from backend.core.logging import get_logger
from backend.models.knowledge_base import KnowledgeBaseDocument
from backend.models.assistant import AssistantConfig
from backend.models.user import User
from backend.schemas.knowledge_base import KnowledgeBaseStatus, KnowledgeBaseDocumentResponse
from backend.services.pinecone_service import PineconeService
from backend.utils.storage import get_file_path, ensure_directory_exists

# Попытка импорта библиотек для обработки различных форматов файлов
try:
    import PyPDF2
    PDF_SUPPORT = True
except ImportError:
    PDF_SUPPORT = False

try:
    import docx
    DOCX_SUPPORT = True
except ImportError:
    DOCX_SUPPORT = False

logger = get_logger(__name__)

# Константы
MAX_CHARS_PER_ASSISTANT = 1000000
MAX_DOCUMENTS_PER_ASSISTANT = 10
ALLOWED_EXTENSIONS = {'txt', 'pdf', 'docx', 'md', 'rtf'}
CHUNK_SIZE = 600  # токенов
CHUNK_OVERLAP = 100  # токенов

class KnowledgeBaseService:
    """Service for working with knowledge base"""
    
    @classmethod
    async def get_knowledge_base_status(cls, db: Session, assistant_id: str) -> KnowledgeBaseStatus:
        """Get knowledge base status for assistant"""
        try:
            # Get documents for assistant
            documents = db.query(KnowledgeBaseDocument).filter(
                KnowledgeBaseDocument.assistant_id == assistant_id
            ).all()
            
            # Calculate total chars
            total_chars = sum(doc.chars_count for doc in documents)
            
            # Create response
            response = KnowledgeBaseStatus(
                total_documents=len(documents),
                total_chars=total_chars,
                max_chars=MAX_CHARS_PER_ASSISTANT,
                max_documents=MAX_DOCUMENTS_PER_ASSISTANT,
                documents=[
                    KnowledgeBaseDocumentResponse(
                        id=str(doc.id),
                        assistant_id=str(doc.assistant_id),
                        filename=doc.filename,
                        original_filename=doc.original_filename,
                        content_type=doc.content_type,
                        size=doc.size,
                        chars_count=doc.chars_count,
                        status=doc.status,
                        processed=doc.processed,
                        error_message=doc.error_message,
                        created_at=doc.created_at,
                        updated_at=doc.updated_at
                    )
                    for doc in documents
                ]
            )
            
            return response
        except Exception as e:
            logger.error(f"Failed to get knowledge base status: {str(e)}")
            raise
    
    @classmethod
    async def upload_document(cls, db: Session, file: UploadFile, assistant_id: str, user_id: str) -> KnowledgeBaseDocumentResponse:
        """Upload document to knowledge base"""
        try:
            # Check if assistant exists and belongs to user
            assistant = db.query(AssistantConfig).filter(
                AssistantConfig.id == assistant_id,
                AssistantConfig.user_id == user_id
            ).first()
            
            if not assistant:
                raise HTTPException(
                    status_code=status.HTTP_404_NOT_FOUND,
                    detail="Assistant not found or doesn't belong to user"
                )
            
            # Check file extension
            extension = file.filename.split('.')[-1].lower() if '.' in file.filename else ''
            if extension not in ALLOWED_EXTENSIONS:
                raise HTTPException(
                    status_code=status.HTTP_400_BAD_REQUEST,
                    detail=f"File type not allowed. Allowed types: {', '.join(ALLOWED_EXTENSIONS)}"
                )
            
            # Check number of documents
            doc_count = db.query(KnowledgeBaseDocument).filter(
                KnowledgeBaseDocument.assistant_id == assistant_id
            ).count()
            
            if doc_count >= MAX_DOCUMENTS_PER_ASSISTANT:
                raise HTTPException(
                    status_code=status.HTTP_400_BAD_REQUEST,
                    detail=f"Maximum number of documents ({MAX_DOCUMENTS_PER_ASSISTANT}) reached"
                )
            
            # Get user's OpenAI API key
            user = db.query(User).filter(User.id == user_id).first()
            if not user or not user.openai_api_key:
                raise HTTPException(
                    status_code=status.HTTP_400_BAD_REQUEST,
                    detail="User has no OpenAI API key"
                )
            
            # Read file content
            content = await file.read()
            file_size = len(content)
            
            # Generate file path
            uploads_dir = ensure_directory_exists("backend/static/uploads")
            file_path = get_file_path(file.filename, user_id, assistant_id, unique=True)
            full_path = os.path.join(uploads_dir, file_path)
            
            # Save file to disk
            os.makedirs(os.path.dirname(full_path), exist_ok=True)
            with open(full_path, "wb") as f:
                f.write(content)
                
            # Get content type
            content_type = file.content_type
            
            # Extract text
            text_content = await cls._extract_text(content, extension, content_type)
            
            # Count characters
            chars_count = len(text_content)
            
            # Check current chars count
            current_chars = sum(
                doc.chars_count for doc in db.query(KnowledgeBaseDocument).filter(
                    KnowledgeBaseDocument.assistant_id == assistant_id
                ).all()
            )
            
            if current_chars + chars_count > MAX_CHARS_PER_ASSISTANT:
                # Remove saved file
                if os.path.exists(full_path):
                    os.remove(full_path)
                    
                raise HTTPException(
                    status_code=status.HTTP_400_BAD_REQUEST,
                    detail=f"Document would exceed maximum character limit ({MAX_CHARS_PER_ASSISTANT})"
                )
            
            # Create document record
            document = KnowledgeBaseDocument(
                assistant_id=assistant_id,
                filename=os.path.splitext(file.filename)[0],
                original_filename=file.filename,
                file_path=file_path,
                content_type=content_type,
                size=file_size,
                chars_count=chars_count,
                status="pending"
            )
            
            db.add(document)
            db.commit()
            db.refresh(document)
            
            # Start processing in background
            asyncio.create_task(cls._process_document(str(document.id), text_content, user.openai_api_key, str(assistant_id)))
            
            # Create response
            response = KnowledgeBaseDocumentResponse(
                id=str(document.id),
                assistant_id=str(document.assistant_id),
                filename=document.filename,
                original_filename=document.original_filename,
                content_type=document.content_type,
                size=document.size,
                chars_count=document.chars_count,
                status=document.status,
                processed=document.processed,
                error_message=document.error_message,
                created_at=document.created_at,
                updated_at=document.updated_at
            )
            
            return response
        except HTTPException:
            raise
        except Exception as e:
            logger.error(f"Failed to upload document: {str(e)}")
            if 'db' in locals() and db.is_active:
                db.rollback()
            raise HTTPException(
                status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
                detail=f"Failed to upload document: {str(e)}"
            )
    
    @classmethod
    async def delete_document(cls, db: Session, document_id: str, assistant_id: str, user_id: str) -> bool:
        """Delete document from knowledge base"""
        try:
            # Check if assistant exists and belongs to user
            assistant = db.query(AssistantConfig).filter(
                AssistantConfig.id == assistant_id,
                AssistantConfig.user_id == user_id
            ).first()
            
            if not assistant:
                raise HTTPException(
                    status_code=status.HTTP_404_NOT_FOUND,
                    detail="Assistant not found or doesn't belong to user"
                )
            
            # Get document
            document = db.query(KnowledgeBaseDocument).filter(
                KnowledgeBaseDocument.id == document_id,
                KnowledgeBaseDocument.assistant_id == assistant_id
            ).first()
            
            if not document:
                raise HTTPException(
                    status_code=status.HTTP_404_NOT_FOUND,
                    detail="Document not found"
                )
            
            # Delete vectors from Pinecone
            await PineconeService.delete_vectors(
                filter={"source": document.original_filename},
                namespace=str(assistant_id)
            )
            
            # Delete file from disk
            file_path = os.path.join("backend/static/uploads", document.file_path)
            if os.path.exists(file_path):
                os.remove(file_path)
            
            # Delete document from database
            db.delete(document)
            db.commit()
            
            return True
        except HTTPException:
            raise
        except Exception as e:
            logger.error(f"Failed to delete document: {str(e)}")
            if 'db' in locals() and db.is_active:
                db.rollback()
            raise HTTPException(
                status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
                detail=f"Failed to delete document: {str(e)}"
            )
    
    @classmethod
    async def sync_knowledge_base(cls, db: Session, assistant_id: str, user_id: str) -> bool:
        """Synchronize knowledge base (reprocess all documents)"""
        try:
            # Check if assistant exists and belongs to user
            assistant = db.query(AssistantConfig).filter(
                AssistantConfig.id == assistant_id,
                AssistantConfig.user_id == user_id
            ).first()
            
            if not assistant:
                raise HTTPException(
                    status_code=status.HTTP_404_NOT_FOUND,
                    detail="Assistant not found or doesn't belong to user"
                )
            
            # Get user's OpenAI API key
            user = db.query(User).filter(User.id == user_id).first()
            if not user or not user.openai_api_key:
                raise HTTPException(
                    status_code=status.HTTP_400_BAD_REQUEST,
                    detail="User has no OpenAI API key"
                )
            
            # Delete all vectors from Pinecone
            await PineconeService.delete_vectors(
                filter={},
                namespace=str(assistant_id),
                delete_all=True
            )
            
            # Get all documents for assistant
            documents = db.query(KnowledgeBaseDocument).filter(
                KnowledgeBaseDocument.assistant_id == assistant_id
            ).all()
            
            # Mark all documents as pending
            for document in documents:
                document.status = "pending"
                document.processed = False
                document.error_message = None
            
            db.commit()
            
            # Process each document
            for document in documents:
                # Read file content
                file_path = os.path.join("backend/static/uploads", document.file_path)
                if not os.path.exists(file_path):
                    document.status = "error"
                    document.error_message = "File not found"
                    continue
                
                with open(file_path, "rb") as f:
                    content = f.read()
                
                # Extract text
                extension = document.original_filename.split('.')[-1].lower() if '.' in document.original_filename else ''
                text_content = await cls._extract_text(content, extension, document.content_type)
                
                # Start processing in background
                asyncio.create_task(cls._process_document(str(document.id), text_content, user.openai_api_key, str(assistant_id)))
            
            db.commit()
            return True
        except HTTPException:
            raise
        except Exception as e:
            logger.error(f"Failed to synchronize knowledge base: {str(e)}")
            if 'db' in locals() and db.is_active:
                db.rollback()
            raise HTTPException(
                status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
                detail=f"Failed to synchronize knowledge base: {str(e)}"
            )
    
    @classmethod
    async def search_knowledge_base(cls, query: str, assistant_id: str, top_k: int = 5, user_api_key: Optional[str] = None) -> List[Dict[str, Any]]:
        """Search knowledge base for query"""
        try:
            # Create embedding for query
            query_embedding = await PineconeService.create_embeddings(query, user_api_key=user_api_key)
            
            # Query Pinecone
            results = await PineconeService.query_vectors(
                query_embedding=query_embedding,
                namespace=assistant_id,
                top_k=top_k
            )
            
            # Extract matches
            matches = []
            for match in results.get("matches", []):
                matches.append({
                    "score": match.get("score", 0),
                    "text": match.get("metadata", {}).get("text", ""),
                    "source": match.get("metadata", {}).get("source", "")
                })
            
            return matches
        except Exception as e:
            logger.error(f"Failed to search knowledge base: {str(e)}")
            raise
    
    @classmethod
    async def _extract_text(cls, content: bytes, extension: str, content_type: str) -> str:
        """Extract text from file content"""
        # Text files
        if extension == 'txt' or content_type == 'text/plain':
            return content.decode('utf-8', errors='ignore')
        
        # Markdown files
        if extension == 'md':
            return content.decode('utf-8', errors='ignore')
        
        # PDF files
        if extension == 'pdf' and PDF_SUPPORT:
            try:
                reader = PyPDF2.PdfReader(io.BytesIO(content))
                text = ""
                for i in range(len(reader.pages)):
                    text += reader.pages[i].extract_text()
                return text
            except Exception as e:
                logger.error(f"Error extracting text from PDF: {str(e)}")
                return "Error extracting text from PDF"
        
        # DOCX files
        if extension == 'docx' and DOCX_SUPPORT:
            try:
                doc = docx.Document(io.BytesIO(content))
                text = ""
                for paragraph in doc.paragraphs:
                    text += paragraph.text + "\n"
                return text
            except Exception as e:
                logger.error(f"Error extracting text from DOCX: {str(e)}")
                return "Error extracting text from DOCX"
        
        # RTF files (simple conversion)
        if extension == 'rtf':
            # Basic RTF stripping
            try:
                text = content.decode('utf-8', errors='ignore')
                # Remove RTF markers
                text = ' '.join(text.split('\\')[1:])
                return text
            except Exception as e:
                logger.error(f"Error extracting text from RTF: {str(e)}")
                return "Error extracting text from RTF"
        
        return "Unsupported file format"
    
    @classmethod
    async def _chunk_text(cls, text: str, chunk_size: int = CHUNK_SIZE, chunk_overlap: int = CHUNK_OVERLAP) -> List[str]:
        """Split text into chunks of specified token size with overlap"""
        try:
            # Use tiktoken for token counting
            enc = tiktoken.get_encoding("cl100k_base")  # Encoder compatible with GPT-4
            
            # Tokenize text
            tokens = enc.encode(text)
            
            # Split into chunks
            chunks = []
            for i in range(0, len(tokens), chunk_size - chunk_overlap):
                # Get chunk_size tokens or remaining tokens if less than chunk_size
                chunk_tokens = tokens[i:i + chunk_size]
                chunk_text = enc.decode(chunk_tokens)
                chunks.append(chunk_text)
            
            return chunks
        except Exception as e:
            logger.error(f"Error chunking text: {str(e)}")
            # Fallback to character-based chunking if tiktoken fails
            chunks = []
            chunk_size_chars = 1000  # approximate chunk size in characters
            overlap_chars = 200  # approximate overlap in characters
            
            for i in range(0, len(text), chunk_size_chars - overlap_chars):
                chunk_text = text[i:i + chunk_size_chars]
                chunks.append(chunk_text)
            
            return chunks
    
    @classmethod
    async def _process_document(cls, document_id: str, text_content: str, api_key: str, assistant_id: str) -> None:
        """Process document and upload vectors to Pinecone"""
        from backend.db.session import SessionLocal
        
        db = SessionLocal()
        try:
            # Get document
            document = db.query(KnowledgeBaseDocument).filter(KnowledgeBaseDocument.id == document_id).first()
            
            if not document:
                logger.error(f"Document not found: {document_id}")
                return
            
            # Update status to processing
            document.status = "processing"
            db.commit()
            
            # Split text into chunks
            chunks = await cls._chunk_text(text_content)
            
            # Create vectors
            vectors = []
            for i, chunk in enumerate(chunks):
                # Create embedding for chunk
                embedding = await PineconeService.create_embeddings(chunk, user_api_key=api_key)
                
                # Create vector
                vector = {
                    "id": f"{document_id}_{i}",
                    "values": embedding,
                    "metadata": {
                        "text": chunk,
                        "source": document.original_filename,
                        "chunk_index": i
                    }
                }
                
                vectors.append(vector)
            
            # Upsert vectors to Pinecone
            success = await PineconeService.upsert_vectors(vectors, namespace=assistant_id)
            
            if success:
                # Update document status
                document.status = "processed"
                document.processed = True
            else:
                # Update document status
                document.status = "error"
                document.error_message = "Failed to upload vectors to Pinecone"
            
            db.commit()
        except Exception as e:
            logger.error(f"Failed to process document: {str(e)}")
            
            # Update document status
            try:
                if db.is_active:
                    document = db.query(KnowledgeBaseDocument).filter(KnowledgeBaseDocument.id == document_id).first()
                    
                    if document:
                        document.status = "error"
                        document.error_message = str(e)
                        db.commit()
            except Exception as db_error:
                logger.error(f"Failed to update document status: {str(db_error)}")
        finally:
            db.close()
